"""Generate client SEO recap document — Guivarche Déménagement, June 2026."""
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent / "Recap_prestations_SEO_Guivarche_Juin_2026.docx"

BRAND = RGBColor(0x1C, 0x39, 0x57)
MUTED = RGBColor(0x64, 0x74, 0x8B)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def style_run(run, *, bold=False, size=11, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = BRAND if level == 1 else RGBColor(0x33, 0x41, 0x55)


def add_para(doc: Document, text: str, *, bold=False, bullet=False):
    if bullet:
        p = doc.add_paragraph(text, style="List Bullet")
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        style_run(run, bold=bold)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p


def add_table(doc: Document, headers: list[str], rows: list[list[str]], col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "E8EEF4")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                style_run(run, bold=True, size=10, color=BRAND)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
            for p in cells[ci].paragraphs:
                for run in p.runs:
                    style_run(run, size=10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # --- Cover ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("GUIVARCHE DÉMÉNAGEMENT")
    style_run(r, bold=True, size=14, color=BRAND)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("Document de synthèse — Prestations SEO")
    style_run(r2, bold=True, size=20, color=BRAND)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run(
        f"Site : guivarche-demenagement.fr\n"
        f"Date d'intervention : 1er juin 2026\n"
        f"Document généré le : {date.today().strftime('%d/%m/%Y')}"
    )
    style_run(r3, size=11, color=MUTED)

    doc.add_paragraph()
    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r4 = intro.add_run(
        "Le présent document récapitule l'ensemble des prestations SEO réalisées sur le site "
        "Guivarche Déménagement : correction des anomalies remontées par l'audit technique "
        "(balises meta, titres dupliqués, attributs alt, structure H1), optimisation des "
        "longueurs pour le référencement naturel, exclusion des landing pages campagne "
        "de l'indexation Google, ainsi que des optimisations UX/SEO complémentaires "
        "(page de confirmation des LP, hiérarchie des titres Hx, liens sortants nofollow, "
        "coordonnées cliquables dans le pied de page)."
    )
    style_run(r4, size=11)

    doc.add_page_break()

    # --- 1. Contexte ---
    add_heading(doc, "1. Contexte et objectifs", 1)
    add_para(
        doc,
        "Le site Guivarche Déménagement (application React, domaine guivarche-demenagement.fr) "
        "présentait plusieurs anomalies SEO identifiées via des outils d'audit (Screaming Frog "
        "et équivalents) : meta descriptions absentes ou dupliquées, titres identiques sur "
        "plusieurs URLs, images sans attribut alt, et balises H1 mal structurées.",
    )
    add_para(doc, "Objectifs de l'intervention :", bold=True)
    add_para(doc, "Corriger chaque anomalie sur les pages indexables du site.", bullet=True)
    add_para(doc, "Centraliser la gestion SEO (titles, descriptions, canonical) dans un module unique.", bullet=True)
    add_para(doc, "Optimiser les longueurs des balises pour limiter la troncature dans les résultats Google.", bullet=True)
    add_para(doc, "Exclure les landing pages campagne (/lp/*) du référencement organique (noindex).", bullet=True)
    add_para(doc, "Améliorer l'expérience post-formulaire et la conformité technique (confirmation LP, structure Hx, nofollow, contact cliquable).", bullet=True)

    # --- 2. Meta descriptions ---
    add_heading(doc, "2. Balises meta description manquantes", 1)
    add_heading(doc, "Problème identifié", 2)
    add_para(
        doc,
        "Plusieurs routes appelaient déjà un hook usePageMeta, mais le module central "
        "pageMeta.ts était absent ou incomplet. Le HTML initial (index.html) et le contenu "
        "pré-rendu réutilisaient la description de la page d'accueil, ce qui produisait "
        "des descriptions manquantes ou dupliquées sur de nombreuses URLs.",
    )
    add_heading(doc, "Actions réalisées", 2)
    add_para(doc, "Création du hook react/src/hooks/usePageMeta.ts : injection dynamique de document.title, balise <meta name=\"description\"> et lien canonical par route.", bullet=True)
    add_para(doc, "Création et enrichissement de react/src/seo/pageMeta.ts : descriptions dédiées pour l'ensemble des pages indexables.", bullet=True)
    add_para(doc, "Mise à jour de react/index.html pour l'accueil (crawl initial avant hydratation React).", bullet=True)

    add_heading(doc, "Pages couvertes", 2)
    add_table(
        doc,
        ["Page", "Route"],
        [
            ["Accueil", "/"],
            ["Tarifs", "/tarif"],
            ["Solutions", "/solution"],
            ["FAQ", "/faq"],
            ["Déménagement particulier", "/demenagement-particulier"],
            ["Déménagement entreprise", "/demenagement-entreprise"],
            ["Blog (index)", "/blog"],
            ["Contact", "/contact"],
            ["RGPD", "/rgpd"],
            ["Mentions légales", "/mentions-legales"],
            ["Île-de-France", "/ile-de-france"],
            ["Déménagement national", "/demenagement-national"],
            ["International", "/international"],
            ["Formules déménagement", "/formules-demenagement"],
            ["Page en construction", "/en-construction"],
            ["Page 404", "/*"],
            ["9 articles de blog", "/blog/{slug}"],
        ],
        col_widths=[6, 8],
    )

    # --- 3. Title dupliquées ---
    add_heading(doc, "3. Balises title dupliquées — titres uniques par page", 1)
    add_heading(doc, "Problème identifié", 2)
    add_para(
        doc,
        "De nombreuses pages affichaient le même title « Guivarche Déménagement » ou une "
        "variante générique, pénalisant le référencement et la différenciation dans les "
        "résultats de recherche.",
    )
    add_heading(doc, "Actions réalisées", 2)
    add_para(doc, "Attribution d'un title SEO unique à chaque page principale, conformément aux briefs client.", bullet=True)
    add_para(doc, "Définition de titres explicites pour les 9 articles de blog dans BLOG_ARTICLE_META.", bullet=True)
    add_para(doc, "Audit global des doublons title/description sur les pages indexables ; réécriture des entrées en conflit.", bullet=True)
    add_para(doc, "Les landing pages campagne (/lp/*) conservent des titres descriptifs mais sont exclues de l'indexation (voir section 8).", bullet=True)

    add_heading(doc, "Exemples de titles assignés (pages principales)", 2)
    add_table(
        doc,
        ["Page", "Meta title"],
        [
            ["Accueil", "Déménageur Paris & IDF : Particuliers et Pro - Guivarche"],
            ["Tarifs", "Tarifs Déménagement : Estimez Votre Budget - Guivarche"],
            ["Solutions", "Nos Solutions de Déménagement sur Mesure - Guivarche"],
            ["FAQ", "FAQ Déménagement : Tarifs, Délais, Devis | Guivarche"],
            ["Particulier", "Service de Déménagement pour Particuliers - Guivarche"],
            ["Entreprise", "Déménagement d'Entreprise: Service Professionnel - Guivarche"],
            ["Blog", "Conseils et Astuces pour un Déménagement Réussi - Guivarche"],
            ["Contact", "Contactez Nos Experts du Déménagement - Guivarche"],
        ],
        col_widths=[4, 10],
    )

    add_heading(doc, "Articles de blog — titres uniques", 2)
    add_table(
        doc,
        ["Article", "Meta title"],
        [
            ["Guide complet", "Guide Complet pour Réussir son Déménagement Facilement"],
            ["Objets fragiles", "Comment Emballer des Objets Fragiles pour un Déménagement"],
            ["Déménagement écologique", "Comment Réussir un Déménagement Écologique - Guivarche"],
            ["Longue distance", "Réussir un Déménagement Longue Distance : Conseils et Astuces"],
            ["Check-list", "Check-list Ultime pour un Déménagement Sans Oublis - Guivarche"],
            ["Avec enfants", "Déménager avec des Enfants : Guide Pratique - Guivarche"],
            ["Erreurs à éviter", "Erreurs à Éviter Lors d'un Déménagement - Guivarche"],
            ["Budget", "Budget Déménagement : Astuces pour Économiser - Guivarche"],
            ["Assurance", "Assurance Déménagement : Guide Complet - Guivarche"],
        ],
        col_widths=[4, 10],
    )

    # --- 4. Alt ---
    add_heading(doc, "4. Balises attribut alt manquantes", 1)
    add_heading(doc, "Problème identifié", 2)
    add_para(
        doc,
        "Plusieurs images du site présentaient un attribut alt vide (alt=\"\") ou absent, "
        "affectant l'accessibilité (WCAG) et le référencement image.",
    )
    add_heading(doc, "Actions réalisées", 2)
    add_para(doc, "Création du fichier central react/src/seo/imageAlts.ts pour centraliser les textes alt validés.", bullet=True)
    add_para(doc, "Fonction getBlogImageAlt() pour mapper automatiquement les images Unsplash des articles de blog.", bullet=True)
    add_para(doc, "Mise à jour des composants et pages concernés.", bullet=True)

    add_heading(doc, "Images corrigées (extrait)", 2)
    add_table(
        doc,
        ["Emplacement", "Texte alt attribué"],
        [
            ["Section IA (ai.jpeg)", "Équipe déménageurs professionnels expérimentés"],
            ["Logos presse — La Revue Durable", "Logo La Revue Durable"],
            ["Logos presse — la provence", "Logo la provence"],
            ["Logos presse — l'express", "Logo l'express"],
            ["Logos presse — midi libre", "Logo midi libre"],
            ["Articles blog (Unsplash)", "Textes descriptifs uniques par photo (ex. déménagement résidentiel, maison moderne…)"],
            ["Listing blog + articles", "imageAlt dédié par article dans blogPosts.ts"],
            ["Composant bento-demo", "Alt unique pour chaque carte visuelle"],
            ["Landing pages (galerie)", "Alt descriptifs pour les photos d'équipe Guivarche"],
        ],
        col_widths=[5, 9],
    )

    add_heading(doc, "Fichiers modifiés", 2)
    add_para(doc, "HomeDesigned.tsx, Blog.tsx, BlogArticle.tsx, blogPosts.ts, bento-demo.tsx, paris.tsx, seine92.tsx, particulier.tsx, pro.tsx.", bullet=False)

    # --- 5. H1 ---
    add_heading(doc, "5. Réécriture des balises H1 en doublon", 1)
    add_heading(doc, "Problème identifié", 2)
    add_para(
        doc,
        "La page d'accueil comportait un H1 peu orienté SEO (nom de marque + slogan dans "
        "le même bloc H1), et la page tarifs utilisait un intitulé générique « Tarification ». "
        "Cette structure nuisait à la clarté sémantique pour les moteurs de recherche.",
    )
    add_heading(doc, "Actions réalisées", 2)
    add_table(
        doc,
        ["Page", "Avant", "Après"],
        [
            [
                "Accueil (/)",
                "H1 « Guivarche Déménagement » + slogan dans le même H1",
                "H1 orienté SEO : « Entreprise de Déménagement pour Particuliers et Professionnels en France. » — slogan déplacé en paragraphe (<p>)",
            ],
            [
                "Tarifs (/tarif)",
                "H1 « Tarification »",
                "H1 « Tarifs de Déménagement : Estimez le Coût de Votre Projet »",
            ],
        ],
        col_widths=[3, 5, 6],
    )
    add_para(
        doc,
        "Structure responsive conservée (versions desktop et mobile) avec un seul H1 principal par vue.",
    )

    # --- 6. Longueurs meta ---
    add_heading(doc, "6. Optimisation des longueurs meta (title & description)", 1)
    add_heading(doc, "Problème identifié", 2)
    add_para(
        doc,
        "Certaines balises étaient trop longues (troncature dans les SERP Google, ~60 caractères "
        "pour le title) ou trop courtes (descriptions insuffisamment informatives, < 120 caractères).",
    )
    add_heading(doc, "Méthodologie", 2)
    add_para(doc, "Audit de 32 entrées dans pageMeta.ts (pages organiques, landing pages, articles blog).", bullet=True)
    add_para(doc, "Cibles retenues : title entre 40 et 65 caractères ; description entre 120 et 165 caractères.", bullet=True)

    add_heading(doc, "Corrections principales", 2)
    add_table(
        doc,
        ["Page / zone", "Problème", "Correction"],
        [
            ["Accueil", "Title/description génériques ou mal calibrés", "Title ~56 car. ; description ~152 car. (Paris, IDF, devis 24 h)"],
            ["Déménagement particulier", "Description trop courte (113 car.)", "Description enrichie ~144 car. avec localisation et CTA"],
            ["Formules déménagement", "Title trop long (68 car.)", "Raccourci : « Formules Déménagement : Éco, Standard, Premium - Guivarche »"],
            ["5 articles blog", "Descriptions trop courtes ou trop longues", "Réécriture entre 146 et 164 caractères"],
            ["7 landing pages", "Meta trop courtes", "Titles et descriptions spécifiques par campagne (Paris, 92, pro, etc.)"],
        ],
        col_widths=[3.5, 4.5, 6],
    )

    # --- 7. FAQ, IDF, National ---
    add_heading(doc, "7. Pages FAQ, Île-de-France et National", 1)
    add_para(
        doc,
        "Suite à contrôle complémentaire, trois pages présentaient encore des titles proches "
        "de la limite de troncature (~61 caractères) et des descriptions perfectibles.",
    )
    add_table(
        doc,
        ["Page", "Route", "Meta title (final)", "Description"],
        [
            [
                "FAQ",
                "/faq",
                "FAQ Déménagement : Tarifs, Délais, Devis | Guivarche",
                "~158 car. — tarifs, délais, assurance, emballage, devis gratuit 24 h",
            ],
            [
                "Île-de-France",
                "/ile-de-france",
                "Déménageur Île-de-France : Paris & Banlieue - Guivarche",
                "~154 car. — Paris, petite couronne, équipes salariées, devis 24 h",
            ],
            [
                "National",
                "/demenagement-national",
                "Déménagement National & Longue Distance - Guivarche",
                "~154 car. — longue distance France, transport sécurisé, suivi personnalisé",
            ],
        ],
        col_widths=[2.5, 3, 5, 4.5],
    )

    # --- 8. LP noindex ---
    add_heading(doc, "8. Landing pages — exclusion du référencement (noindex)", 1)
    add_heading(doc, "Exigence", 2)
    add_para(
        doc,
        "Les pages /lp/* (campagnes Google Ads et acquisition payante) ne doivent pas "
        "apparaître dans l'index Google ni concurrencer les pages organiques du site.",
    )
    add_heading(doc, "Mécanismes déployés (défense en profondeur)", 2)
    add_table(
        doc,
        ["Couche", "Fichier / composant", "Action"],
        [
            ["Meta robots + canonical", "usePageMeta.ts, landingPageMeta()", "robots: noindex, follow ; suppression du lien canonical sur les LP"],
            ["Navigation SPA", "ScrollToTop.tsx, isLandingPath()", "Synchronisation du meta robots à chaque changement de route"],
            ["robots.txt", "public/robots.txt", "Disallow: /lp/ pour Googlebot, Bingbot, Twitterbot, facebookexternalhit et User-agent *"],
            ["En-tête HTTP", "public/.htaccess", "X-Robots-Tag: noindex, follow pour /lp/ et /tunnel/"],
            ["Sitemap", "public/sitemap.xml", "Aucune URL /lp/* listée"],
            ["Prerender", "scripts/prerender.mjs", "Routes /lp/* exclues du pré-rendu SEO"],
        ],
        col_widths=[3, 4.5, 6.5],
    )

    add_heading(doc, "Landing pages couvertes", 2)
    add_table(
        doc,
        ["Route", "Usage"],
        [
            ["/lp/paris", "Campagne Paris"],
            ["/lp/hauts-de-seine", "Campagne Hauts-de-Seine (92)"],
            ["/lp/pro", "Campagne professionnels"],
            ["/lp/particulier", "Campagne particuliers"],
            ["/lp/demenagement-entreprise", "Campagne déménagement entreprise"],
            ["/lp/demenagement-particulier", "Campagne déménagement particulier"],
            ["/lp/ile-de-france", "Campagne Île-de-France"],
        ],
        col_widths=[6, 8],
    )
    add_para(
        doc,
        "Toute future URL sous le préfixe /lp/ est automatiquement traitée en noindex, "
        "sans modification supplémentaire.",
    )

    # --- 9. Prestations complémentaires ---
    add_heading(doc, "9. Prestations complémentaires (UX, technique & SEO)", 1)

    add_heading(doc, "9.1 Page de confirmation pour les landing pages", 2)
    add_heading(doc, "Contexte", 3)
    add_para(
        doc,
        "Les landing pages campagne (/lp/paris, /lp/pro, /lp/particulier, /lp/hauts-de-seine, etc.) "
        "proposent un formulaire de devis intégré. Auparavant, la confirmation s'affichait sur la "
        "même page après envoi, sans parcours unifié avec le tunnel de devis principal.",
    )
    add_heading(doc, "Actions réalisées", 3)
    add_para(
        doc,
        "Après soumission réussie du formulaire (composant DevisForm partagé), l'utilisateur est "
        "redirigé vers la route /tunnel/devis/confirmation, identique au tunnel organique.",
        bullet=True,
    )
    add_para(
        doc,
        "Suppression de l'affichage inline « Demande envoyée » sur les LP pour éviter les doublons "
        "de message et harmoniser l'expérience.",
        bullet=True,
    )
    add_para(
        doc,
        "La route de confirmation reste accessible sans blocage RouteGuard pour les visiteurs "
        "provenant des campagnes payantes.",
        bullet=True,
    )
    add_heading(doc, "Fichiers concernés", 3)
    add_para(doc, "react/src/pages/paris.tsx (DevisForm) — redirection navigate() vers /tunnel/devis/confirmation.", bullet=True)
    add_para(doc, "react/src/App.tsx — route quote-confirmation déjà en place.", bullet=True)

    add_heading(doc, "9.2 Vérification et correction de la structure Hx (H1–H6)", 2)
    add_heading(doc, "Contexte", 3)
    add_para(
        doc,
        "Un audit de la hiérarchie des titres (outil type Heading Map / inspection manuelle) "
        "avait révélé des anomalies sur plusieurs templates : H1 absents ou multiples, sauts de "
        "niveau (H1 → H3), doublons de H2 (sections avis desktop/mobile), libellés « Sophie » "
        "en H2, placeholders Figma en H3, etc.",
    )
    add_heading(doc, "Actions réalisées (site entier)", 3)
    add_para(doc, "Tunnel de devis (App.tsx) : un H1 par étape (formulaire, méthodes, sr-only sur les autres étapes) ; « Mes étapes » en H2 ; profils Sophie en <p>.", bullet=True)
    add_para(doc, "Accueil (HomeDesigned.tsx) : H1 sr-only unique ; titres visuels hero en <p> ; réordonnancement H2/H3 section « 3 méthodes » ; suppression des faux H3 lorem.", bullet=True)
    add_para(doc, "Pages services, zones, blog, FAQ, solution, formules : H1 hero conservé ou ajouté (motion.h1) ; sections en H2 → H3.", bullet=True)
    add_para(doc, "BlogArticle.tsx : H1 article ; sidebar « Table des matières » / « Tags » en H3 ; sections contenu en H2/H3.", bullet=True)
    add_para(doc, "Landing pages : hero H1 avant les H2 de section ; pitch hero en <p> si doublon avec H1.", bullet=True)
    add_para(doc, "Footer.tsx : libellés de colonnes passés en <p> (éviter des H3 orphelins sur toutes les pages).", bullet=True)
    add_para(doc, "Composant GoogleReviewsSection.tsx : un seul H2 + widget Elfsight (remplace les blocs dupliqués desktop/mobile).", bullet=True)
    add_para(doc, "Pages tunnel info / options / devis : H2 « prestations incluses » ; un seul H1 sur l'étape devis.", bullet=True)

    add_heading(doc, "Périmètre pages contrôlées", 3)
    add_table(
        doc,
        ["Zone", "Exemples de routes"],
        [
            ["Site principal", "/, /solution, /faq, /tarif, /formules-demenagement, /contact"],
            ["Offres", "/demenagement-particulier, /demenagement-entreprise, /ile-de-france, /demenagement-national, /international"],
            ["Blog", "/blog, /blog/{slug}"],
            ["Tunnel", "/tunnel/mes-coordonnees, étapes volume, /tunnel/devis/confirmation"],
            ["Landing pages", "/lp/paris, /lp/pro, /lp/particulier, /lp/hauts-de-seine, /lp/demenagement-*"],
        ],
        col_widths=[4, 10],
    )

    add_heading(doc, "9.3 Attribut rel=\"nofollow\" sur les liens sortants", 2)
    add_heading(doc, "Contexte", 3)
    add_para(
        doc,
        "Les liens vers des domaines tiers (réseaux sociaux, Trustpilot, Pages Jaunes, Google "
        "Avis, WhatsApp, etc.) doivent transmettre rel=\"nofollow\" (avec noopener noreferrer "
        "lorsque target=\"_blank\") pour limiter la dilution du PageRank et sécuriser l'ouverture "
        "dans un nouvel onglet.",
    )
    add_heading(doc, "Actions réalisées", 3)
    add_para(doc, "Création de react/src/constants/externalLink.ts : constante EXTERNAL_LINK_REL = \"nofollow noopener noreferrer\".", bullet=True)
    add_para(doc, "Application sur le pied de page : Trustpilot, Pages Jaunes, Facebook, Instagram, LinkedIn, YouTube, TikTok, Reddit, Pinterest, Tumblr.", bullet=True)
    add_para(doc, "Application sur la page Contact (cartes réseaux sociaux + Pages Jaunes).", bullet=True)
    add_para(doc, "Application sur les LP (lien avis Google), partage blog (Pages Jaunes), bouton WhatsApp flottant (StickyContactButtons).", bullet=True)
    add_para(doc, "Composant get-in-touch.tsx : nofollow automatique pour les href https:// externes.", bullet=True)
    add_para(
        doc,
        "Note : les liens générés dynamiquement par le widget Elfsight (avis Google) ne sont pas "
        "contrôlables dans le code source ; seuls les liens statiques du site sont couverts.",
    )

    add_heading(doc, "9.4 Coordonnées cliquables dans le pied de page", 2)
    add_heading(doc, "Contexte", 3)
    add_para(
        doc,
        "L'e-mail et les numéros de téléphone du footer étaient affichés en texte brut (<p>), "
        "sans action mailto: ni tel: sur mobile et desktop.",
    )
    add_heading(doc, "Actions réalisées", 3)
    add_table(
        doc,
        ["Coordonnée", "Lien", "Affichage"],
        [
            ["E-mail", "mailto:contact@guivarche-demenagement.fr", "contact@guivarche-demenagement.fr"],
            ["Mobile", "tel:+33746326678", "+33 7 46 32 66 78"],
            ["Fixe", "tel:+33 1 89 70 33 24", "+33 1 89 70 33 24 (aligné page Contact)"],
        ],
        col_widths=[3, 5.5, 5.5],
    )
    add_para(
        doc,
        "Styles hover et focus clavier ajoutés pour l'accessibilité. L'adresse postale reste en texte simple.",
    )
    add_heading(doc, "Fichier modifié", 3)
    add_para(doc, "react/src/components/layout/Footer.tsx", bullet=True)

    # --- 10. Synthèse ---
    add_heading(doc, "10. Synthèse technique et recommandations", 1)
    add_heading(doc, "Fichiers créés ou modifiés (principaux)", 2)
    add_table(
        doc,
        ["Fichier", "Rôle"],
        [
            ["react/src/hooks/usePageMeta.ts", "Gestion dynamique title, description, canonical, robots"],
            ["react/src/seo/pageMeta.ts", "Registre central des meta + LP noindex"],
            ["react/src/seo/imageAlts.ts", "Textes alt centralisés"],
            ["react/src/components/ScrollToTop.tsx", "Robots meta synchronisé à la navigation"],
            ["react/index.html", "Meta initiales page d'accueil"],
            ["react/public/robots.txt", "Blocage crawl /lp/"],
            ["react/public/sitemap.xml", "Sitemap pages indexables uniquement"],
            ["react/public/.htaccess", "X-Robots-Tag serveur"],
            ["react/scripts/prerender.mjs", "Exclusion LP du prerender"],
            ["Pages React (HomeDesigned, Tarification, Blog*, landing*)", "H1, alt, branchement usePageMeta"],
            ["react/src/constants/externalLink.ts", "rel nofollow sur liens sortants"],
            ["react/src/components/GoogleReviewsSection.tsx", "Section avis unique (structure Hx)"],
            ["react/src/components/layout/Footer.tsx", "Coordonnées mailto/tel cliquables"],
            ["react/src/pages/paris.tsx (DevisForm)", "Redirection confirmation LP"],
            ["App.tsx, BlogArticle.tsx, pages services/zones/LP", "Corrections hiérarchie Hx"],
        ],
        col_widths=[5.5, 8.5],
    )

    add_heading(doc, "Actions recommandées post-déploiement", 2)
    add_para(doc, "Exécuter npm run build et le script de prerender pour mettre à jour le dossier dist/.", bullet=True)
    add_para(doc, "Vérifier dans Google Search Console (Inspection d'URL) qu'une landing page /lp/paris affiche bien « Exclue : noindex ».", bullet=True)
    add_para(doc, "Contrôler que les pages organiques (/faq, /ile-de-france, etc.) présentent title et description uniques dans les SERP après réindexation.", bullet=True)
    add_para(doc, "Tester un envoi de formulaire depuis /lp/paris et vérifier l'arrivée sur /tunnel/devis/confirmation.", bullet=True)
    add_para(doc, "Recontrôler la hiérarchie Hx (extension Heading Map ou équivalent) sur les pages clés après déploiement.", bullet=True)

    add_heading(doc, "Bilan des prestations", 2)
    add_para(
        doc,
        "L'intervention couvre l'ensemble du périmètre SEO technique demandé : correction "
        "des meta descriptions manquantes, élimination des titles dupliqués, complétion des "
        "attributs alt, restructuration des H1, calibrage des longueurs meta sur 32 pages, "
        "optimisation ciblée FAQ / Île-de-France / National, verrouillage noindex des "
        "landing pages campagne sur quatre niveaux (meta, SPA, robots.txt, serveur), ainsi que "
        "les prestations complémentaires : redirection vers une page de confirmation unifiée "
        "pour les LP, audit et correction de la structure Hx sur tout le site, ajout de "
        "rel=\"nofollow\" sur les liens sortants statiques, et coordonnées cliquables (e-mail "
        "et téléphones) dans le footer.",
    )

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rf = footer.add_run(
        f"Document généré le {date.today().strftime('%d/%m/%Y')} — Confidentiel — Guivarche Déménagement"
    )
    style_run(rf, size=9, color=MUTED)

    return doc


def main():
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Document généré : {OUTPUT}")


if __name__ == "__main__":
    main()
